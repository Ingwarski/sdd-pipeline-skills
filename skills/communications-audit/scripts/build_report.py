#!/usr/bin/env python3
"""Validate Communications Audit JSON and build a polished editable DOCX."""

from __future__ import annotations

import argparse
import json
import math
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CANONICAL_DIMENSIONS = [
    ("communicator", "Communicator", "Who are we?"),
    ("audience", "Audience", "Who are we addressing?"),
    ("desired_effect", "Desired Effect", "What should happen?"),
    ("message", "Message", "What are we saying?"),
    ("channel", "Channel", "Where is it communicated?"),
    ("timing", "Timing", "When does it reach the audience?"),
    ("execution", "Execution", "How is it communicated?"),
]

SEVERITIES = ("Critical", "Major", "Moderate", "Minor")
LEVELS = ("High", "Medium", "Low")
PRIORITIES = ("Now", "Next", "Later")
EVIDENCE_CLASSES = ("Observation", "Inference", "Hypothesis", "External fact")

NAVY = "0B2545"
BLUE = "2E74B5"
CHARCOAL = "252B33"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "1F5D42"
RULE = "D7DBE2"
FONT = "Arial"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


class InputError(ValueError):
    pass


def require_dict(value: Any, path: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise InputError(f"{path} must be an object")
    return value


def require_list(value: Any, path: str) -> list[Any]:
    if not isinstance(value, list):
        raise InputError(f"{path} must be an array")
    return value


def require_text(container: dict[str, Any], key: str, path: str) -> str:
    value = container.get(key)
    if not isinstance(value, str) or not value.strip():
        raise InputError(f"{path}.{key} must be non-empty text")
    return value.strip()


def optional_text(container: dict[str, Any], key: str, default: str = "") -> str:
    value = container.get(key, default)
    return value.strip() if isinstance(value, str) else default


def validate_choice(value: Any, allowed: tuple[str, ...], path: str) -> str:
    if value not in allowed:
        raise InputError(f"{path} must be one of: {', '.join(allowed)}")
    return str(value)


def validate_report(data: Any) -> dict[str, Any]:
    report = require_dict(data, "report")
    metadata = require_dict(report.get("metadata"), "metadata")
    for key in ("client", "subject", "date", "auditor", "confidentiality", "language", "scope"):
        require_text(metadata, key, "metadata")
    accent = optional_text(metadata, "accent_color")
    if accent and not re.fullmatch(r"#?[0-9A-Fa-f]{6}", accent):
        raise InputError("metadata.accent_color must be a six-digit hex color")
    logo_path = optional_text(metadata, "auditor_logo_path")
    if logo_path and not Path(logo_path).expanduser().is_file():
        raise InputError(f"metadata.auditor_logo_path not found: {logo_path}")

    executive = require_dict(report.get("executive_summary"), "executive_summary")
    require_text(executive, "headline", "executive_summary")
    require_text(executive, "conclusion", "executive_summary")
    for key in ("strengths", "risks", "priorities"):
        values = require_list(executive.get(key), f"executive_summary.{key}")
        if not values or any(not isinstance(item, str) or not item.strip() for item in values):
            raise InputError(f"executive_summary.{key} must contain non-empty text items")

    scorecard = require_list(report.get("scorecard"), "scorecard")
    ids = [item.get("id") if isinstance(item, dict) else None for item in scorecard]
    expected_ids = [item[0] for item in CANONICAL_DIMENSIONS]
    if ids != expected_ids:
        raise InputError(f"scorecard IDs must appear exactly in this order: {', '.join(expected_ids)}")

    applicable_weight_total = 0.0
    original_weight_total = 0.0
    timing_applicable = True
    for index, item_any in enumerate(scorecard):
        item = require_dict(item_any, f"scorecard[{index}]")
        dimension_id = expected_ids[index]
        applicable = item.get("applicable", True)
        if not isinstance(applicable, bool):
            raise InputError(f"scorecard[{index}].applicable must be true or false")
        if not applicable and dimension_id != "timing":
            raise InputError("Only Timing may be marked not applicable")
        weight = item.get("weight")
        if not isinstance(weight, (int, float)) or isinstance(weight, bool) or weight <= 0:
            raise InputError(f"scorecard[{index}].weight must be a positive number")
        original_weight_total += float(weight)
        if applicable:
            score = item.get("score")
            if not isinstance(score, int) or isinstance(score, bool) or not 1 <= score <= 5:
                raise InputError(f"scorecard[{index}].score must be an integer from 1 to 5")
            validate_choice(item.get("confidence"), LEVELS, f"scorecard[{index}].confidence")
            require_text(item, "rationale", f"scorecard[{index}]")
            applicable_weight_total += float(weight)
        else:
            reason = optional_text(item, "rationale")
            if not reason:
                raise InputError("A not-applicable Timing dimension requires a rationale")
        if dimension_id == "timing":
            timing_applicable = applicable

    if applicable_weight_total <= 0:
        raise InputError("At least one scorecard dimension must be applicable")
    if not math.isclose(original_weight_total, 100.0, abs_tol=0.01):
        raise InputError(f"Original scorecard weights must total 100%; got {original_weight_total:g}%")

    current_context = require_dict(report.get("current_context"), "current_context")
    context_applicable = current_context.get("applicable")
    if not isinstance(context_applicable, bool):
        raise InputError("current_context.applicable must be true or false")
    if timing_applicable:
        if not context_applicable:
            raise InputError("Applicable Timing requires current_context.applicable=true")
        require_text(current_context, "as_of", "current_context")
        require_text(current_context, "summary", "current_context")
        context_sources = require_list(current_context.get("sources"), "current_context.sources")
        if not context_sources or any(not isinstance(item, str) or not item.strip() for item in context_sources):
            raise InputError("Applicable Timing requires at least one current-context source ID")
    elif context_applicable:
        raise InputError("current_context cannot be applicable when Timing is marked N/A")

    findings = require_list(report.get("findings"), "findings")
    if not findings:
        raise InputError("findings must contain at least one finding")
    grave_count = 0
    finding_source_refs: list[tuple[int, list[str]]] = []
    for index, item_any in enumerate(findings):
        item = require_dict(item_any, f"findings[{index}]")
        for key in ("title", "evidence", "implication", "recommendation", "owner", "success_metric"):
            require_text(item, key, f"findings[{index}]")
        severity = validate_choice(item.get("severity"), SEVERITIES, f"findings[{index}].severity")
        if severity in ("Critical", "Major"):
            grave_count += 1
        evidence_class = validate_choice(item.get("evidence_class"), EVIDENCE_CLASSES, f"findings[{index}].evidence_class")
        validate_choice(item.get("impact"), LEVELS, f"findings[{index}].impact")
        validate_choice(item.get("effort"), LEVELS, f"findings[{index}].effort")
        validate_choice(item.get("confidence"), LEVELS, f"findings[{index}].confidence")
        validate_choice(item.get("priority"), PRIORITIES, f"findings[{index}].priority")
        steps = require_list(item.get("steps"), f"findings[{index}].steps")
        canonical_labels = {item[1] for item in CANONICAL_DIMENSIONS}
        if not steps or any(step not in canonical_labels for step in steps):
            raise InputError(f"findings[{index}].steps must use canonical dimension labels")
        source_refs = item.get("source_ids", [])
        source_refs = require_list(source_refs, f"findings[{index}].source_ids")
        if any(not isinstance(source_id, str) or not source_id.strip() for source_id in source_refs):
            raise InputError(f"findings[{index}].source_ids must contain source IDs")
        if evidence_class == "External fact" and not source_refs:
            raise InputError(f"findings[{index}] uses External fact and requires source_ids")
        finding_source_refs.append((index, source_refs))
        for image_key in ("image_path", "image_caption"):
            if image_key in item and item[image_key] is not None and not isinstance(item[image_key], str):
                raise InputError(f"findings[{index}].{image_key} must be text")

    allowed_findings = max(7, grave_count)
    if len(findings) > allowed_findings:
        raise InputError(
            f"{len(findings)} findings exceed the allowed {allowed_findings}; consolidate lower-severity items or retain only all Critical/Major findings"
        )

    roadmap = require_dict(report.get("roadmap"), "roadmap")
    for horizon in ("now", "next", "later"):
        items = require_list(roadmap.get(horizon), f"roadmap.{horizon}")
        for index, item_any in enumerate(items):
            item = require_dict(item_any, f"roadmap.{horizon}[{index}]")
            for key in ("action", "owner", "success_metric"):
                require_text(item, key, f"roadmap.{horizon}[{index}]")

    measurement = require_list(report.get("measurement"), "measurement")
    if not 3 <= len(measurement) <= 5:
        raise InputError("measurement must contain 3 to 5 KPIs")
    for index, item_any in enumerate(measurement):
        item = require_dict(item_any, f"measurement[{index}]")
        for key in ("metric", "definition", "data_source", "cadence", "owner", "baseline", "target"):
            require_text(item, key, f"measurement[{index}]")

    benchmarking = require_dict(report.get("benchmarking"), "benchmarking")
    requested = benchmarking.get("requested")
    if not isinstance(requested, bool):
        raise InputError("benchmarking.requested must be true or false")
    comparators = require_list(benchmarking.get("comparators"), "benchmarking.comparators")
    if requested:
        require_text(benchmarking, "summary", "benchmarking")
        if not comparators:
            raise InputError("Requested benchmarking requires at least one comparator")
    elif comparators:
        raise InputError("benchmarking.comparators must be empty when benchmarking was not requested")

    assumptions = require_list(report.get("assumptions"), "assumptions")
    specialist = require_list(report.get("specialist_followups"), "specialist_followups")
    for path, values in (("assumptions", assumptions), ("specialist_followups", specialist)):
        if any(not isinstance(item, str) or not item.strip() for item in values):
            raise InputError(f"{path} must contain text items")

    sources = require_list(report.get("sources"), "sources")
    source_ids: set[str] = set()
    for index, item_any in enumerate(sources):
        item = require_dict(item_any, f"sources[{index}]")
        source_id = require_text(item, "id", f"sources[{index}]")
        if source_id in source_ids:
            raise InputError(f"Duplicate source ID: {source_id}")
        source_ids.add(source_id)
        require_text(item, "title", f"sources[{index}]")
        if optional_text(item, "url") and not re.match(r"^https?://", item["url"].strip()):
            raise InputError(f"sources[{index}].url must start with http:// or https://")
    if timing_applicable:
        missing = [source_id for source_id in current_context["sources"] if source_id not in source_ids]
        if missing:
            raise InputError(f"current_context references undefined source IDs: {', '.join(missing)}")
    for finding_index, refs in finding_source_refs:
        missing = [source_id for source_id in refs if source_id not in source_ids]
        if missing:
            raise InputError(f"findings[{finding_index}] references undefined source IDs: {', '.join(missing)}")
    if requested and not sources:
        raise InputError("Requested benchmarking requires external sources")

    return report


def score_report(scorecard: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], int, str]:
    total = sum(float(item["weight"]) for item in scorecard if item.get("applicable", True))
    scored: list[dict[str, Any]] = []
    raw_overall = 0.0
    for item in scorecard:
        entry = dict(item)
        if item.get("applicable", True):
            effective_weight = float(item["weight"]) / total * 100.0
            normalized = (int(item["score"]) - 1) / 4.0
            raw_overall += effective_weight * normalized
            entry["effective_weight"] = effective_weight
        else:
            entry["effective_weight"] = 0.0
        scored.append(entry)
    overall = int(math.floor(raw_overall + 0.5))
    if overall <= 24:
        band = "Critical"
    elif overall <= 49:
        band = "Weak"
    elif overall <= 69:
        band = "Functional"
    elif overall <= 84:
        band = "Strong"
    else:
        band = "Excellent"
    return scored, overall, band


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size: float = 11, color: str = CHARCOAL, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = RULE, size: int = 4, style: str = "single") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA; got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_bottom_border(paragraph, color: str = BLUE, size: int = 14) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def style_paragraph(paragraph, before: float = 0, after: float = 6, line: float = 1.1, keep_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep_next:
        fmt.keep_with_next = True


def replace_cell_text(cell, text: str, *, bold: bool = False, color: str = CHARCOAL, size: float = 9.2,
                      align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    style_paragraph(paragraph, after=0, line=1.05)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_label_paragraph(doc, label: str, text: str, *, color: str = CHARCOAL, after: float = 5) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=after, line=1.1)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=NAVY, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=10.5, color=color)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        style_paragraph(paragraph, after=4, line=1.08)
        run = paragraph.add_run(item.strip())
        set_run_font(run, size=10.5)


def add_heading(doc, text: str, level: int = 1, *, page_break_before: bool = False) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 11.5}[level],
        color={1: NAVY, 2: BLUE, 3: NAVY}[level],
        bold=True,
    )


def add_source_line(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="Source Text")
    style_paragraph(paragraph, before=2, after=4, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED)


def add_table_header(table, labels: list[str]) -> None:
    row = table.rows[0]
    set_repeat_header(row)
    for index, label in enumerate(labels):
        set_cell_shading(row.cells[index], NAVY)
        replace_cell_text(row.cells[index], label, bold=True, color=WHITE, size=8.8)


def configure_styles(doc) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.08

    if "Source Text" not in styles:
        style = styles.add_style("Source Text", 1)
    else:
        style = styles["Source Text"]
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(8.5)
    style.font.color.rgb = rgb(MUTED)


def configure_document(doc, metadata: dict[str, Any]) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(paragraph, after=0, line=1.0)
    left = paragraph.add_run(f"{metadata['client']}  •  {metadata['date']}    ")
    set_run_font(left, size=8.5, color=MUTED)
    add_page_field(paragraph)

    props = doc.core_properties
    props.title = f"Communications Audit — {metadata['subject']}"
    props.subject = "Evidence-led sales and marketing communications audit"
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.keywords = ""


def add_cover(doc, metadata: dict[str, Any], executive: dict[str, Any]) -> None:
    logo_path = optional_text(metadata, "auditor_logo_path")
    if logo_path:
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(logo, before=10, after=16)
        inline = logo.add_run().add_picture(str(Path(logo_path).expanduser()), width=Inches(1.25))
        inline._inline.docPr.set("descr", optional_text(metadata, "auditor_logo_alt", "Auditor logo"))
    else:
        spacer = doc.add_paragraph()
        style_paragraph(spacer, before=28, after=16)

    kicker = doc.add_paragraph()
    style_paragraph(kicker, after=10)
    run = kicker.add_run("COMMUNICATIONS AUDIT")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    style_paragraph(title, after=8, line=1.0, keep_next=True)
    run = title.add_run(metadata["subject"])
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=22, line=1.05)
    run = subtitle.add_run(f"Prepared for {metadata['client']}")
    set_run_font(run, size=14, color=CHARCOAL)
    add_bottom_border(subtitle, color=BLUE, size=12)

    rows = [
        ("Auditor", metadata["auditor"]),
        ("Date", metadata["date"]),
        ("Scope", metadata["scope"]),
        ("Language", metadata["language"]),
        ("Status", metadata["confidentiality"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for index, (label, value) in enumerate(rows):
        replace_cell_text(table.rows[index].cells[0], label.upper(), bold=True, color=MUTED, size=8.5)
        replace_cell_text(table.rows[index].cells[1], value, color=CHARCOAL, size=10.5)
    set_table_geometry(table, [1700, 7660])
    set_table_borders(table, color=WHITE, size=0, style="nil")

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=26, after=6)
    run = paragraph.add_run("EXECUTIVE THESIS")
    set_run_font(run, size=8.5, color=BLUE, bold=True)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    set_table_borders(callout, color=PALE_BLUE, size=6)
    set_cell_shading(callout.cell(0, 0), PALE_BLUE)
    replace_cell_text(callout.cell(0, 0), executive["headline"], bold=True, color=NAVY, size=13)

def add_score_panel(doc, overall: int, band: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2200, 7160])
    set_table_borders(table, color=PALE_BLUE, size=6)
    set_cell_shading(table.cell(0, 0), NAVY)
    set_cell_shading(table.cell(0, 1), PALE_BLUE)

    score_p = table.cell(0, 0).paragraphs[0]
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(score_p, after=0, line=1.0)
    score_run = score_p.add_run(str(overall))
    set_run_font(score_run, size=28, color=WHITE, bold=True)
    label_run = score_p.add_run("\n/ 100")
    set_run_font(label_run, size=9, color=WHITE, bold=True)

    body_p = table.cell(0, 1).paragraphs[0]
    style_paragraph(body_p, after=0, line=1.05)
    band_run = body_p.add_run(f"{band} communications effectiveness")
    set_run_font(band_run, size=13, color=NAVY, bold=True)
    detail_run = body_p.add_run("\nWeighted from the seven applicable dimensions; confidence is reported separately.")
    set_run_font(detail_run, size=9.5, color=CHARCOAL)


def add_executive_summary(doc, executive: dict[str, Any], overall: int, band: str) -> None:
    add_heading(doc, executive["headline"], 1, page_break_before=True)
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=12, line=1.15)
    run = paragraph.add_run(executive["conclusion"])
    set_run_font(run, size=11.5, color=CHARCOAL)
    add_score_panel(doc, overall, band)

    add_heading(doc, "What already works", 2)
    add_bullets(doc, executive["strengths"])
    add_heading(doc, "What puts performance at risk", 2)
    add_bullets(doc, executive["risks"])
    add_heading(doc, "What leadership should do first", 2)
    add_bullets(doc, executive["priorities"])
def add_scorecard(doc, scorecard: list[dict[str, Any]], current_context: dict[str, Any]) -> None:
    add_heading(
        doc,
        "The communication is strongest where the seven dimensions align",
        1,
        page_break_before=True,
    )
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=10)
    run = paragraph.add_run(
        "Ratings use a 1–5 evidence rubric. Effective weights are recalculated when Timing is genuinely not applicable."
    )
    set_run_font(run, size=10.5, color=CHARCOAL)

    table = doc.add_table(rows=1, cols=5)
    add_table_header(table, ["Dimension", "Score", "Weight", "Confidence", "Evidence-led rationale"])
    canonical = {item[0]: item[1:] for item in CANONICAL_DIMENSIONS}
    for item in scorecard:
        row = table.add_row()
        label, question = canonical[item["id"]]
        score_text = str(item["score"]) if item.get("applicable", True) else "N/A"
        weight_text = f"{item['effective_weight']:.1f}%" if item.get("applicable", True) else "0%"
        confidence = item.get("confidence", "—") if item.get("applicable", True) else "—"
        values = [f"{label}\n{question}", score_text, weight_text, confidence, item["rationale"]]
        for index, value in enumerate(values):
            replace_cell_text(
                row.cells[index],
                value,
                bold=index == 0,
                color=NAVY if index == 0 else CHARCOAL,
                size=8.6 if index == 4 else 8.8,
                align=WD_ALIGN_PARAGRAPH.CENTER if index in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT,
            )
        if not item.get("applicable", True):
            for cell in row.cells:
                set_cell_shading(cell, LIGHT)
    set_table_geometry(table, [1750, 750, 900, 1400, 4560])
    set_table_borders(table)
    add_source_line(doc, "Scoring formula: Σ effective weight × (rating − 1) ÷ 4. Final score is rounded once.")

    if current_context.get("applicable"):
        add_heading(doc, "Current context changes how Timing should be judged", 2)
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, after=4)
        run = paragraph.add_run(current_context["summary"])
        set_run_font(run, size=10.5)
        add_source_line(
            doc,
            f"Current-context research as of {current_context['as_of']} • Sources: {', '.join(current_context['sources'])}",
        )
def severity_color(severity: str) -> str:
    return {"Critical": RED, "Major": GOLD, "Moderate": BLUE, "Minor": MUTED}[severity]


def add_before_after(doc, before: str, after: str) -> None:
    table = doc.add_table(rows=2, cols=2)
    add_table_header(table, ["Before", "Recommended direction"])
    replace_cell_text(table.cell(1, 0), before, color=CHARCOAL, size=9.2)
    replace_cell_text(table.cell(1, 1), after, color=NAVY, size=9.2)
    set_cell_shading(table.cell(1, 1), PALE_BLUE)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)


def add_finding_image(doc, item: dict[str, Any]) -> None:
    image_path = optional_text(item, "image_path")
    if not image_path:
        return
    path = Path(image_path).expanduser()
    if not path.is_file():
        raise InputError(f"Evidence image not found: {path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(paragraph, before=6, after=2)
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(6.15))
    inline._inline.docPr.set("descr", optional_text(item, "image_caption", "Evidence from audited material"))
    caption = optional_text(item, "image_caption")
    if caption:
        add_source_line(doc, caption)


def add_findings(doc, findings: list[dict[str, Any]]) -> None:
    add_heading(doc, "Priority findings translate evidence into action", 1, page_break_before=True)
    for index, item in enumerate(findings, start=1):
        heading = doc.add_paragraph(style="Heading 2")
        heading.paragraph_format.keep_with_next = True
        number_run = heading.add_run(f"{index}. ")
        set_run_font(number_run, size=13, color=BLUE, bold=True)
        title_run = heading.add_run(item["title"])
        set_run_font(title_run, size=13, color=NAVY, bold=True)

        marker = doc.add_paragraph()
        style_paragraph(marker, after=5, line=1.0, keep_next=True)
        severity_run = marker.add_run(item["severity"].upper())
        set_run_font(severity_run, size=8.5, color=severity_color(item["severity"]), bold=True)
        steps_run = marker.add_run("   •   " + " / ".join(item["steps"]))
        set_run_font(steps_run, size=8.5, color=MUTED, bold=True)

        add_label_paragraph(doc, f"Evidence — {item['evidence_class']}", item["evidence"])
        if item.get("source_ids"):
            add_source_line(doc, f"Sources: {', '.join(item['source_ids'])}")
        add_label_paragraph(doc, "Why it matters", item["implication"])
        add_label_paragraph(doc, "Recommendation", item["recommendation"], after=7)

        table = doc.add_table(rows=4, cols=2)
        metadata_rows = [
            ("Priority", item["priority"]),
            ("Impact / effort / confidence", f"{item['impact']} / {item['effort']} / {item['confidence']}"),
            ("Likely owner", item["owner"]),
            ("Success metric", item["success_metric"]),
        ]
        for row_index, (label, value) in enumerate(metadata_rows):
            set_cell_shading(table.rows[row_index].cells[0], LIGHT)
            replace_cell_text(table.rows[row_index].cells[0], label, bold=True, color=NAVY, size=8.7)
            replace_cell_text(table.rows[row_index].cells[1], value, color=CHARCOAL, size=9.0)
        set_table_geometry(table, [2500, 6860])
        set_table_borders(table)

        before = optional_text(item, "before")
        after = optional_text(item, "after")
        if before or after:
            add_before_after(doc, before or "Not supplied", after or "Not supplied")
        add_finding_image(doc, item)

        rule = doc.add_paragraph()
        style_paragraph(rule, before=5, after=5)
        add_bottom_border(rule, color=RULE, size=4)


def add_roadmap(doc, roadmap: dict[str, Any]) -> None:
    add_heading(
        doc,
        "Sequence the work by impact, effort, confidence, and dependency",
        1,
        page_break_before=True,
    )
    table = doc.add_table(rows=1, cols=4)
    add_table_header(table, ["Horizon", "Action", "Owner", "Success metric"])
    horizon_labels = (("now", "Now"), ("next", "Next"), ("later", "Later"))
    for key, label in horizon_labels:
        for item in roadmap[key]:
            row = table.add_row()
            values = [label, item["action"], item["owner"], item["success_metric"]]
            for index, value in enumerate(values):
                replace_cell_text(row.cells[index], value, bold=index == 0, color=NAVY if index == 0 else CHARCOAL, size=8.8)
            if key == "now":
                set_cell_shading(row.cells[0], PALE_BLUE)
            elif key == "later":
                set_cell_shading(row.cells[0], LIGHT)
    set_table_geometry(table, [900, 4260, 1500, 2700])
    set_table_borders(table)


def add_measurement(doc, measurement: list[dict[str, Any]]) -> None:
    add_heading(doc, "Measure whether the communication improves the intended effect", 1)
    table = doc.add_table(rows=1, cols=5)
    add_table_header(table, ["KPI", "Definition", "Source / cadence", "Owner", "Baseline / target"])
    for item in measurement:
        row = table.add_row()
        values = [
            item["metric"],
            item["definition"],
            f"{item['data_source']}\n{item['cadence']}",
            item["owner"],
            f"{item['baseline']}\n{item['target']}",
        ]
        for index, value in enumerate(values):
            replace_cell_text(row.cells[index], value, bold=index == 0, color=NAVY if index == 0 else CHARCOAL, size=8.3)
    set_table_geometry(table, [1500, 2500, 1800, 1400, 2160])
    set_table_borders(table)


def add_benchmarking(doc, benchmarking: dict[str, Any], *, page_break_before: bool = False) -> None:
    if not benchmarking.get("requested"):
        return
    add_heading(
        doc,
        "Requested benchmarking adds an external comparison lens",
        1,
        page_break_before=page_break_before,
    )
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=8)
    run = paragraph.add_run(benchmarking["summary"])
    set_run_font(run, size=10.5)
    table = doc.add_table(rows=1, cols=2)
    add_table_header(table, ["Comparator", "Evidence and implication"])
    for comparator in benchmarking["comparators"]:
        if isinstance(comparator, dict):
            name = optional_text(comparator, "name", "Comparator")
            detail = optional_text(comparator, "detail", optional_text(comparator, "summary", ""))
        else:
            name = str(comparator)
            detail = ""
        row = table.add_row()
        replace_cell_text(row.cells[0], name, bold=True, color=NAVY, size=9)
        replace_cell_text(row.cells[1], detail, color=CHARCOAL, size=9)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table)


def add_method_sources(doc, report: dict[str, Any], *, page_break_before: bool = False) -> None:
    add_heading(
        doc,
        "Method, assumptions, and source record",
        1,
        page_break_before=page_break_before,
    )
    add_heading(doc, "Method", 2)
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=6)
    run = paragraph.add_run(
        "The audit applies seven communication dimensions—Communicator, Audience, Desired Effect, Message, Channel, Timing, and Execution—then prioritizes findings by business impact, effort, confidence, and dependency. The overall score is a transparent weighted normalization of the 1–5 dimension ratings."
    )
    set_run_font(run, size=10.5)

    add_heading(doc, "Assumptions and limitations", 2)
    if report["assumptions"]:
        add_bullets(doc, report["assumptions"])
    else:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, after=4)
        run = paragraph.add_run("No material assumptions beyond the stated audit scope.")
        set_run_font(run, size=10.5)

    if report["specialist_followups"]:
        add_heading(doc, "Specialist follow-ups outside this score", 2)
        add_bullets(doc, report["specialist_followups"])

    add_heading(doc, "Sources", 2)
    if not report["sources"]:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, after=4)
        run = paragraph.add_run("No external sources were used. Findings are based on the supplied audit material.")
        set_run_font(run, size=9.5, color=MUTED)
        return
    for source in report["sources"]:
        parts = [f"[{source['id']}] {source['title']}"]
        publisher = optional_text(source, "publisher")
        published = optional_text(source, "published")
        accessed = optional_text(source, "accessed")
        url = optional_text(source, "url")
        if publisher:
            parts.append(publisher)
        if published:
            parts.append(f"published {published}")
        if accessed:
            parts.append(f"accessed {accessed}")
        if url:
            parts.append(url)
        add_source_line(doc, " • ".join(parts))


def build_report(report: dict[str, Any], output_path: Path) -> tuple[int, str]:
    global BLUE
    accent = optional_text(report["metadata"], "accent_color")
    if accent:
        BLUE = accent.lstrip("#").upper()
    scored, overall, band = score_report(report["scorecard"])
    doc = Document()
    configure_styles(doc)
    configure_document(doc, report["metadata"])
    add_cover(doc, report["metadata"], report["executive_summary"])
    add_executive_summary(doc, report["executive_summary"], overall, band)
    add_scorecard(doc, scored, report["current_context"])
    add_findings(doc, report["findings"])
    add_roadmap(doc, report["roadmap"])
    add_measurement(doc, report["measurement"])
    benchmarking_requested = bool(report["benchmarking"].get("requested"))
    add_benchmarking(doc, report["benchmarking"], page_break_before=benchmarking_requested)
    add_method_sources(doc, report, page_break_before=not benchmarking_requested)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return overall, band


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="UTF-8 report JSON")
    parser.add_argument("--output", type=Path, help="Output DOCX path")
    parser.add_argument("--validate-only", action="store_true", help="Validate JSON without building DOCX")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.input.is_file():
        print(f"ERROR: input file not found: {args.input}", file=sys.stderr)
        return 2
    if not args.validate_only and args.output is None:
        print("ERROR: --output is required unless --validate-only is used", file=sys.stderr)
        return 2
    try:
        data = json.loads(args.input.read_text(encoding="utf-8-sig"))
        report = validate_report(data)
        scored, overall, band = score_report(report["scorecard"])
        if args.validate_only:
            print(f"VALID: overall score {overall}/100 ({band}); {len(report['findings'])} findings")
            return 0
        overall, band = build_report(report, args.output)
        print(f"CREATED: {args.output}")
        print(f"SCORE: {overall}/100 ({band})")
        return 0
    except (InputError, json.JSONDecodeError, OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
